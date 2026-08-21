"""
Canonical resume-extraction logic for the Blue Collar Portal.

This module is the only place that calls Gemini. The API upload endpoint
(server.py) and the batch CLI (scripts/pipeline.py) both go through
extract_candidate(), so the prompt, response schema, retry policy, skill
taxonomy and embedding format stay identical no matter which path a resume
arrives on. Previously those two paths had diverged: they used different
Gemini SDKs, different prompts and different embedding text, which meant
vectors produced by the API were not comparable to vectors produced by the
batch job even though both were written to the same pgvector column.

The Gemini client and the MiniLM model are created lazily on first use so
that importing this module never requires an API key or a model download.
"""

from __future__ import annotations

import io
import logging
import time
from pathlib import Path
from typing import List, Optional

import docx
import httpx
from google import genai
from google.genai import errors as genai_errors
from google.genai import types
from pydantic import BaseModel, Field
from sentence_transformers import SentenceTransformer

from config import GEMINI_API_KEY, GEMINI_MODEL

logger = logging.getLogger(__name__)

EMBEDDING_MODEL = "all-MiniLM-L6-v2"
EMBEDDING_DIMENSIONS = 384
MAX_RETRIES = 6

# Extensions we can hand to Gemini as raw multimodal bytes.
MIME_TYPES = {
    ".pdf": "application/pdf",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
}
# Extensions we convert to plain text locally before sending.
TEXT_EXTENSIONS = {".docx", ".txt"}
SUPPORTED_EXTENSIONS = tuple(MIME_TYPES) + tuple(TEXT_EXTENSIONS)

# Statuses and HTTP codes that represent a transient condition worth retrying.
RETRYABLE_STATUSES = frozenset({
    "RESOURCE_EXHAUSTED",  # quota or rate limit
    "UNAVAILABLE",         # backend temporarily down
    "DEADLINE_EXCEEDED",   # server-side timeout
    "INTERNAL",            # transient server fault
})
RETRYABLE_CODES = frozenset({429, 500, 502, 503, 504})

# Controlled vocabulary the search UI filters on. Free-text skills coming back
# from the model are collapsed onto these before they reach the database.
TAXONOMY = [
    "Electrician", "Plumber", "Carpenter", "Welder", "HVAC Technician",
    "Mason/Bricklayer", "Construction Laborer", "Auto Mechanic", "Diesel Mechanic",
    "Driver (Heavy/Light Vehicle)", "Machine/Plant Operator", "Engineering/Lab Technician",
    "Fabricator/Metal Worker", "Fitter", "Cook/Chef", "Food Preparation/Service",
    "Housekeeping/Cleaner", "Security Guard", "Gardener/Landscaper",
    "General Technician/Helper", "Healthcare Support", "Social/Clinical Worker", "Other",
]
# Bucket used when nothing in the taxonomy fits; not a real trade.
FALLBACK_SKILL = "Other"
_TAXONOMY_LOOKUP = {label.lower(): label for label in TAXONOMY}

PROMPT = f"""Extract resume fields strictly adhering to the provided schema.
Map candidate skills to one or more of these standard categories: {TAXONOMY}.
If no category fits perfectly, use ["{FALLBACK_SKILL}"]."""


class ExtractionError(Exception):
    """
    A resume could not be turned into a candidate record.

    Extraction failure is signalled by raising, never by returning a record with
    empty fields. A resume genuinely containing no phone number and a resume the
    API refused to process must not produce the same value, or a caller writing
    to the database cannot tell a real anonymous candidate from a swallowed
    failure. Carries the filename so a batch caller can report which file failed.

    args must hold (filename, message) rather than one pre-formatted string.
    Celery reconstructs an exception from its result backend as cls(*exc.args),
    so a two-argument __init__ paired with a single-element args cannot be
    rebuilt -- it raises TypeError and Celery substitutes
    UnpickleableExceptionWrapper. That is not cosmetic: it meant /status
    reported every failure under a meaningless type instead of
    TransientExtractionError or PermanentExtractionError, so the whole typed
    hierarchy was invisible to API clients. Found by load testing, not by the
    unit tests, which never round-tripped the exception.
    """

    def __init__(self, filename: str, message: str):
        self.filename = filename
        self.message = message
        super().__init__(filename, message)

    def __str__(self) -> str:
        return f"{self.filename}: {self.message}"


class TransientExtractionError(ExtractionError):
    """Failure that may succeed on a later attempt: 5xx, rate limit, network fault."""


class PermanentExtractionError(ExtractionError):
    """Failure retrying cannot fix: unsupported type, unreadable file, bad request."""


class CandidateSchema(BaseModel):
    name: Optional[str] = Field(description="Candidate's full legal name")
    raw_title: Optional[str] = Field(description="Primary job title or headline")
    experience_years: Optional[int] = Field(description="Total years of experience as an integer")
    phone: Optional[str] = Field(description="Contact phone number")
    email: Optional[str] = Field(description="Contact email address")
    skills: List[str] = Field(description="List of matched skill categories from taxonomy")


_client: Optional[genai.Client] = None
_embedder: Optional[SentenceTransformer] = None


def get_client() -> genai.Client:
    """Gemini client, created on first use. Requires GEMINI_API_KEY."""
    global _client
    if _client is None:
        if not GEMINI_API_KEY:
            raise RuntimeError(
                "GEMINI_API_KEY is not set. Copy .env.example to .env and fill it in."
            )
        _client = genai.Client(api_key=GEMINI_API_KEY)
    return _client


def get_embedder() -> SentenceTransformer:
    """MiniLM encoder, loaded on first use (downloads once, then cached)."""
    global _embedder
    if _embedder is None:
        logger.info("Loading embedding model %s...", EMBEDDING_MODEL)
        _embedder = SentenceTransformer(EMBEDDING_MODEL)
    return _embedder


def is_retryable(exc: BaseException) -> bool:
    """
    Decide whether a failed Gemini call is worth another attempt.

    Classified on the SDK's typed exceptions and HTTP status codes rather than
    by substring-matching the message. Message text is not part of the SDK's
    contract and changes between releases, and matching on "503" would also
    misfire on any error string that happens to contain those digits.

    Retryable: network faults that never reached the API, the SDK's 5xx
    ServerError family, and rate-limit/quota responses. Not retryable: an
    invalid API key, a malformed file, or a bad request -- waiting does not
    fix any of those.
    """
    # Never reached the API; always worth another attempt.
    if isinstance(exc, httpx.TransportError):
        return True
    if isinstance(exc, genai_errors.APIError):
        if isinstance(exc, genai_errors.ServerError):
            return True
        if getattr(exc, "code", None) in RETRYABLE_CODES:
            return True
        return getattr(exc, "status", None) in RETRYABLE_STATUSES
    return False


def normalize_skills(skills: Optional[List[str]]) -> List[str]:
    """
    Collapse model-returned skills onto the taxonomy.

    The prompt already asks for taxonomy labels, but the model does drift into
    free text ("MIG Welding" instead of "Welder"), so this is enforced rather
    than trusted. Unrecognised labels are dropped; if nothing survives we fall
    back to ["Other"] so the column is never empty.
    """
    if not skills:
        return [FALLBACK_SKILL]

    normalized: List[str] = []
    for skill in skills:
        if not isinstance(skill, str):
            continue
        canonical = _TAXONOMY_LOOKUP.get(skill.strip().lower())
        if canonical and canonical not in normalized:
            normalized.append(canonical)

    return normalized or [FALLBACK_SKILL]


def build_embedding_text(record: dict) -> str:
    """
    Text fed to MiniLM.

    Every path that writes to the embedding column must call this. Two callers
    building this string independently is exactly how the Level 2 bug happened:
    both were individually correct and silently disagreed with each other,
    placing candidates in different regions of the same vector space.
    """
    skills_text = ", ".join(record.get("skills") or []) or "None"
    return (
        f"Title: {record.get('raw_title')}. "
        f"Experience: {record.get('experience_years')} years. "
        f"Skills: {skills_text}."
    )


def embed_text(text: str) -> List[float]:
    """Encode arbitrary text with MiniLM. Used for search queries too."""
    return get_embedder().encode(text).tolist()


def _docx_to_text(file_bytes: bytes) -> str:
    """
    Flatten a .docx into plain text.

    Table cells are included because resumes frequently lay contact details or
    skills out in a borderless table, and reading only paragraphs silently
    drops them.
    """
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(part for part in parts if part.strip())


def _build_contents(file_bytes: bytes, extension: str) -> list:
    """
    Turn a file into the `contents` payload for Gemini.

    Text-native formats are decoded locally rather than uploaded as bytes:
    Gemini cannot read .docx directly, and sending .txt as a binary part wastes
    tokens for no gain.
    """
    if extension == ".docx":
        return [PROMPT, _docx_to_text(file_bytes)]
    if extension == ".txt":
        return [PROMPT, file_bytes.decode("utf-8", errors="replace")]
    return [
        PROMPT,
        types.Part.from_bytes(data=file_bytes, mime_type=MIME_TYPES[extension]),
    ]


def extract_candidate(filename: str, file_bytes: bytes) -> dict:
    """
    Extract structured candidate data from a resume and embed it. One attempt.

    Takes the filename (used for format detection and provenance) and the raw
    bytes, so callers holding an upload in memory never need to touch disk.
    Returns the candidate dict plus a 384-dim "embedding" list, ready to insert
    into the candidates table.

    Raises TransientExtractionError or PermanentExtractionError on failure and
    never returns a partial record, so a caller cannot mistake a failure for a
    successful extraction of a sparse resume.

    Retry policy deliberately lives with the caller rather than here. A Celery
    worker must yield its slot with self.retry(countdown=...); blocking it with
    time.sleep would hold a concurrency slot open for the whole backoff and
    defeat the point of the queue. The CLI can afford to block, and uses
    extract_candidate_with_backoff below.
    """
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise PermanentExtractionError(
            filename,
            f"unsupported file type {extension!r}; "
            f"supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}",
        )

    try:
        contents = _build_contents(file_bytes, extension)
    except Exception as exc:
        # A corrupt .docx or undecodable payload will not read correctly on a
        # second attempt either.
        raise PermanentExtractionError(filename, f"could not read file: {exc}") from exc

    try:
        response = get_client().models.generate_content(
            model=GEMINI_MODEL,
            contents=contents,
            config=types.GenerateContentConfig(
                temperature=0,
                response_mime_type="application/json",
                response_schema=CandidateSchema,
            ),
        )
    except Exception as exc:
        if is_retryable(exc):
            raise TransientExtractionError(filename, f"transient API failure: {exc}") from exc
        raise PermanentExtractionError(filename, f"API rejected the request: {exc}") from exc

    try:
        record = CandidateSchema.model_validate_json(response.text).model_dump()
    except Exception as exc:
        # Schema-constrained output that still failed to parse is a bad response,
        # not a transport problem; retrying the identical request rarely helps.
        raise PermanentExtractionError(
            filename, f"model returned unparseable JSON: {exc}"
        ) from exc

    record["source_file"] = filename
    record["skills"] = normalize_skills(record.get("skills"))
    record["embedding"] = embed_text(build_embedding_text(record))
    return record


def extract_candidate_with_backoff(
    filename: str, file_bytes: bytes, max_attempts: int = MAX_RETRIES
) -> dict:
    """
    Blocking-retry wrapper around extract_candidate, for the CLI only.

    Retries TransientExtractionError with exponential backoff; permanent
    failures propagate immediately. Must not be used inside a Celery task --
    see the note in extract_candidate.
    """
    last_attempt = max_attempts - 1
    for attempt in range(max_attempts):
        try:
            return extract_candidate(filename, file_bytes)
        except TransientExtractionError as exc:
            if attempt == last_attempt:
                logger.error(
                    "Gave up on %s after %d attempts: %s", filename, max_attempts, exc
                )
                raise
            wait_time = (2 ** attempt) + 2
            logger.warning(
                "Transient failure on %s, retrying in %ss (attempt %d/%d): %s",
                filename, wait_time, attempt + 1, max_attempts, exc,
            )
            time.sleep(wait_time)
